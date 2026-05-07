# report_generator.py
import os
import shutil
import datetime
import json
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from .utils_text import parse_json_insight_to_table
from .utils_analysis import wasted_spend_analyzer
from .utils_text import clean, safe_parse_gemini_json

# Default images source (import from your storage utils if available)
try:
    from audit.storage_utils import REPORT_IMAGES_DIR
except Exception:
    REPORT_IMAGES_DIR = os.path.join(os.path.dirname(__file__), "report_images")


# ---------------- helper utilities ---------------------------------------
def add_industry_benchmark_overlay(df, benchmarks):
    """Add delta and flag columns to a dataframe for a few metrics."""
    if df is None or df.empty:
        return df
    df = df.copy()
    for metric in ["CTR", "Avg CPC", "CPA ($)", "CVR"]:
        if metric in df.columns and metric in benchmarks:
            try:
                df[f"{metric} Δ"] = df[metric].astype(float) - float(benchmarks[metric])
            except Exception:
                # If casting fails, skip numeric delta
                df[f"{metric} Δ"] = ""
            def _flag(x):
                try:
                    delta = float(x.get(f"{metric} Δ", 0))
                    base = float(benchmarks[metric])
                    if base == 0:
                        return ""
                    if delta > 0.2 * base:
                        return "Above"
                    if delta < -0.2 * base:
                        return "Below"
                    return ""
                except Exception:
                    return ""
            df[f"{metric} Flag"] = df.apply(_flag, axis=1)
    return df


def _safe_pct_display(val):
    """Return percentage string. Accepts fraction (0.03) or percent (3)."""
    try:
        n = float(val)
        if 0 <= n <= 1:
            return f"{n * 100:.2f}%"
        return f"{n:.2f}%"
    except Exception:
        return str(val)


def _safe_currency_display(val):
    try:
        return f"${float(val):.2f}"
    except Exception:
        return str(val)


def _ensure_dir(path):
    os.makedirs(path, exist_ok=True)
    return path


# ---------------- main generator -----------------------------------------
def generate_report(
    df_30,
    kw_df,
    hour_pivot,
    hour_raw_df,
    insight_30,
    insight_kw,
    insight_hour,
    geo_df,
    insight_geo,
    wasted_flags,
    wasted_insight,
    lp_audit_rows,
    risk_opp_insights,
    lp_flags=None,
    competitor_df=None,
    output_dir: str = None,
    images_source_dir: str = None,
):
    """
    Generate a Google Ads audit .docx file and return absolute path.

    - output_dir: directory where the .docx (and its _images folder) will be saved.
                  Defaults to /tmp (Cloud Run friendly).
    - images_source_dir: folder where heatmap PNGs live (defaults to REPORT_IMAGES_DIR).
    """
    # Decide output dir
    if output_dir:
        base_out = os.path.abspath(output_dir)
    else:
        base_out = os.environ.get("REPORT_OUTPUT_DIR", "/tmp")
    _ensure_dir(base_out)

    images_source_dir = images_source_dir or REPORT_IMAGES_DIR

    doc = Document()
    doc.add_heading("Google Ads Audit Report", 0)

    def autofit_table(table):
        table.allow_autofit = True
        table.autofit = True
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8)

    def add_table(title, df, columns):
        if df is None or df.empty:
            return
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        autofit_table(table)
        hdr = table.rows[0].cells
        for i, col in enumerate(columns):
            hdr[i].text = col
        for _, row in df.iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(columns):
                val = row.get(col, "")
                if col in ["Avg CPC", "CPA ($)", "Cost ($)"]:
                    cells[i].text = _safe_currency_display(val)
                elif col == "CTR":
                    cells[i].text = _safe_pct_display(val)
                else:
                    try:
                        cells[i].text = "" if pd.isna(val) else str(val)
                    except ValueError:
                        # Happens if val is a list or array
                        cells[i].text = str(val)

    def add_json_insight_section(title, json_text):
        doc.add_heading(title, level=1)
        try:
            df = parse_json_insight_to_table(json_text)
            if df is not None and not df.empty:
                add_table(title, df, df.columns.tolist())
                return
        except Exception:
            pass

        doc.add_paragraph("⚠️ Unable to parse structured insights — showing raw output below.")
        if isinstance(json_text, str):
            doc.add_paragraph(json_text.strip())
        else:
            try:
                doc.add_paragraph(json.dumps(json_text, indent=2))
            except Exception:
                doc.add_paragraph(str(json_text))

    def add_hourly_pivot(pivot):
        doc.add_heading("Hourly Performance Pivot", level=1)
        try:
            day_order = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday", "Saturday", "Sunday"]
            pivot = pivot.reindex(day_order)
        except Exception:
            pass
        try:
            top_levels = getattr(pivot.columns, "levels", None) and pivot.columns.levels[0] or pivot.columns
            for metric in top_levels:
                try:
                    sub_df = pivot[metric].replace("", 0)
                except Exception:
                    sub_df = pivot
                sub_df = sub_df.loc[(sub_df != 0).any(axis=1), (sub_df != 0).any(axis=0)]
                if sub_df.empty:
                    continue
                doc.add_paragraph(f"{metric}")
                table = doc.add_table(rows=sub_df.shape[0] + 1, cols=sub_df.shape[1] + 1)
                table.style = "Table Grid"
                autofit_table(table)
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Day/Hour"
                for j, col in enumerate(sub_df.columns):
                    hdr_cells[j + 1].text = str(col)
                for i, idx in enumerate(sub_df.index):
                    row_cells = table.rows[i + 1].cells
                    row_cells[0].text = str(idx)
                    for j, col in enumerate(sub_df.columns):
                        val = sub_df.loc[idx, col]
                        row_cells[j + 1].text = f"{val:.2f}" if isinstance(val, (int, float)) and val != 0 else ""
        except Exception:
            pass

    def add_heatmaps_and_copy(img_out_dir):
        _ensure_dir(img_out_dir)
        added_any = False
        for metric in ["Clicks", "Conversions", "CVR"]:
            src = os.path.join(images_source_dir, f"{metric}_heatmap.png")
            if not os.path.exists(src):
                src = os.path.join(images_source_dir, f"{metric.lower()}_heatmap.png")
            if os.path.exists(src):
                try:
                    dest = os.path.join(img_out_dir, os.path.basename(src))
                    shutil.copy2(src, dest)
                    doc.add_picture(dest, width=Inches(6))
                    added_any = True
                except Exception:
                    pass
        return added_any

    def add_risks_opportunities():
        try:
            data = risk_opp_insights
            if isinstance(data, str):
                data = json.loads(data)
            if not isinstance(data, dict):
                raise ValueError("Invalid risk_opp_insights format")

            risks = pd.DataFrame(data.get("Risks", []))
            opps = pd.DataFrame(data.get("Opportunities", []))

            doc.add_heading("⚠️ Risks", level=1)
            if not risks.empty:
                add_table("Risks", risks, risks.columns.tolist())
            else:
                doc.add_paragraph("No Risk insights generated.")

            doc.add_heading("✅ Opportunities", level=1)
            if not opps.empty:
                add_table("Opportunities", opps, opps.columns.tolist())
            else:
                doc.add_paragraph("No Opportunity insights generated.")

        except Exception as e:
            doc.add_heading("⚠️ Risks", level=1)
            doc.add_paragraph(f"Failed to parse Risks: {e}")
            doc.add_paragraph(str(risk_opp_insights))

            doc.add_heading("✅ Opportunities", level=1)
            doc.add_paragraph("Parsing failed.")

    # ---------------- build content --------------------------------------
    benchmarks = {"CTR": 0.03, "CVR": 0.05, "Avg CPC": 2.0, "CPA ($)": 20.0}
    try:
        df_30_bench = add_industry_benchmark_overlay(df_30.copy() if df_30 is not None else pd.DataFrame(), benchmarks)
    except Exception:
        df_30_bench = df_30.copy() if df_30 is not None else pd.DataFrame()

    add_table(
        "Campaign Performance (w/ Benchmark Overlay)",
        df_30_bench,
        [
            "Campaign Name",
            "CTR",
            "Cost ($)",
            "Clicks",
            "Impressions",
            "CTR Flag",
            "Avg CPC",
            "Conversions",
            "CPA ($)",
            "Conversion Rate",
        ],
    )
    add_json_insight_section("Campaign Insights", insight_30)

    add_table(
        "Keyword Performance",
        kw_df,
        ["Ad Group", "Keyword", "Match Type", "Quality Score", "Impressions", "Clicks", "CTR", "Avg CPC", "CPA ($)"],
    )
    add_json_insight_section("Keyword Insights", insight_kw)

    if wasted_flags:
        if wasted_insight and isinstance(wasted_insight, str) and wasted_insight.strip():
            add_json_insight_section("Wasted Spend Insights", wasted_insight)
        else:
            doc.add_heading("Wasted Spend Insights", level=1)
            doc.add_paragraph("No insights generated by Gemini.")

    if lp_audit_rows:
        doc.add_heading("Landing Page Audit Insights", level=1)
        for raw_json in lp_audit_rows:
            try:
                data = json.loads(raw_json)
                if isinstance(data, dict):
                    data = [data]
                df = pd.DataFrame(data)
                if df.empty:
                    raise ValueError("Empty LP audit data")
                cols = ["URL"] + [c for c in df.columns if c != "URL"]
                url = df["URL"].iloc[0] if "URL" in df.columns and not df["URL"].empty else "Unknown"
                doc.add_heading(f"Landing Page: {url}", level=2)
                add_table("Landing Page Insights", df, cols)
            except Exception:
                doc.add_paragraph(raw_json.strip() if raw_json else "")

    add_table(
        "Geographical Performance",
        geo_df,
        ["City", "Region", "Country", "Type", "Impressions", "Clicks", "Conversions", "Cost ($)", "CVR", "CPA ($)"],
    )
    add_json_insight_section("Geographical Insights", insight_geo)

    add_hourly_pivot(hour_pivot)
    add_json_insight_section("Hourly Patterns Insights", insight_hour)

    # Prepare output paths and images folder
    filename = f"google_ads_audit_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    out_path = os.path.join(base_out, filename)
    images_out_dir = os.path.splitext(out_path)[0] + "_images"
    _ensure_dir(images_out_dir)

    # Copy/Embed heatmaps if available
    try:
        add_heatmaps_and_copy(images_out_dir)
    except Exception:
        pass

    # Risks & Opportunities
    add_risks_opportunities()

    # Competitor Insights
    if competitor_df is not None and not competitor_df.empty:
        doc.add_heading("🏁 Competitor Insights", level=1)
        cols = competitor_df.columns.tolist()
        add_table("Competitor Insights", competitor_df, cols)
    else:
        doc.add_heading("🏁 Competitor Insights", level=1)
        doc.add_paragraph(
            "Competitor analysis unavailable for this run. "
            "Ensure GEMINI_API_KEY is set and search_term_view returned data. "
            "For richer competitor ad data, set SERPER_API_KEY (free at serper.dev — 2500 searches/month)."
        )

    # Save document
    doc.save(out_path)
    print(f"[SUCCESS] Report saved as {out_path}")

    return os.path.abspath(out_path)