import os
import sys
import io

# Force UTF-8 on Windows to prevent UnicodeEncodeError with emojis
if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import re
import io
from dotenv import load_dotenv

# Load .env file
load_dotenv()
print("DEBUG: .env loaded.")

# Set Google Cloud Credentials if available
SIGNER_KEY_PATH = os.environ.get("REPORT_SIGNER_SECRET_PATH", "campaign-signed-report.json")
if os.path.exists(SIGNER_KEY_PATH):
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.path.abspath(SIGNER_KEY_PATH)
    print(f"DEBUG: GOOGLE_APPLICATION_CREDENTIALS set to {SIGNER_KEY_PATH}")
else:
    print("DEBUG: No GCS credentials found. GCS features may hang or fail.")

# Allow insecure transport for local development only
_redirect = os.environ.get("REDIRECT_URI", "http://localhost:5000/callback")
if _redirect.startswith("http://"):
    os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"
os.environ["OAUTHLIB_RELAX_TOKEN_SCOPE"] = "1"
print(f"DEBUG: OAUTHLIB config set (insecure={'1' if 'OAUTHLIB_INSECURE_TRANSPORT' in os.environ else '0'}).")

import yaml
import json
import ast
from copy import deepcopy
from datetime import timedelta
from datetime import datetime
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
    session,
    send_file,
)
print("DEBUG: flask and basic libs imported.")
from werkzeug.utils import secure_filename
from flask_session import Session
from werkzeug.middleware.proxy_fix import ProxyFix  # ✅ fix HTTPS detection
print("DEBUG: werkzeug/flask_session imported.")
from google_auth_oauthlib.flow import Flow
from google.oauth2 import id_token, service_account
from google.auth.transport import requests
from google.cloud import storage
print("DEBUG: google cloud libs imported.")
from audit.main_runner import generate_google_ads_report  # keep as-is; ensure import works in runtime
print("DEBUG: audit.main_runner imported.")
from audit.storage_utils import upload_file_to_gcs, upload_report_images_to_gcs, REPORT_IMAGES_DIR
print("DEBUG: audit.storage_utils imported.")

# ---- Dev-only: allow insecure oauth only when explicitly requested ----
if os.environ.get("ALLOW_INSECURE_OAUTH", "0") == "1":
    os.environ["OAUTHLIB_INSECURE_TRANSPORT"] = "1"

from audit.chat_with_data import chat_bp


# -------------------- Helpers & config -----------------------------------
def resource_path(relative_path):
    # Use directory of this file as base (works in container & dev)
    if getattr(sys, "frozen", False):
        base_path = os.path.dirname(sys.executable)
    else:
        base_path = os.path.dirname(os.path.abspath(__file__))
    return os.path.join(base_path, relative_path)


template_dir = resource_path("templates")
static_dir = resource_path("static")

app = Flask(__name__, template_folder=template_dir, static_folder=static_dir)

# Register chat blueprint after Flask app is created
app.register_blueprint(chat_bp, url_prefix="/chat")
# Fix proxy headers so URL generation sees correct scheme/host when behind Cloud Run
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# Use environment variables with /tmp defaults for Cloud Run
app.config["UPLOAD_FOLDER"] = os.environ.get("UPLOAD_FOLDER", "/tmp/generated_reports")
app.config["SESSION_TYPE"] = "filesystem"
app.secret_key = os.environ.get("FLASK_SECRET", "mushy_bear")  # replace in prod
app.config["SESSION_FILE_DIR"] = os.environ.get("SESSION_FILE_DIR", "/tmp/flask_session")
app.config["SESSION_PERMANENT"] = False
app.config["SESSION_USE_SIGNER"] = True
Session(app)

USER_TOKENS_DIR = os.environ.get("USER_TOKENS_DIR", "/tmp/user_tokens")
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["SESSION_FILE_DIR"], exist_ok=True)
os.makedirs(REPORT_IMAGES_DIR, exist_ok=True)
os.makedirs(USER_TOKENS_DIR, exist_ok=True)


# ----- GCS (optional persistence) ---------------------------------------
BUCKET_NAME = os.environ.get("BUCKET_NAME")
print(f"DEBUG: BUCKET_NAME is {BUCKET_NAME}")
storage_client = None
bucket = None
if BUCKET_NAME:
    try:
        print("DEBUG: Initializing GCS client...")
        storage_client = storage.Client()
        bucket = storage_client.bucket(BUCKET_NAME)
        print("DEBUG: GCS bucket initialized.")
    except Exception as e:
        print(f"DEBUG: GCS Init failed (likely no creds): {e}")


def download_gcs_to_local(blob_name, local_path):
    if not bucket:
        raise RuntimeError("BUCKET_NAME not set or bucket not initialized")
    blob = bucket.blob(blob_name)
    blob.download_to_filename(local_path)
    return local_path


# Signer secret: prefer secret-mounted path or fallback to local file
SIGNER_KEY_PATH = os.environ.get("REPORT_SIGNER_SECRET_PATH", "campaign-signed-report.json")
_signer_creds = None
if os.path.exists(SIGNER_KEY_PATH):
    try:
        signer_credentials = service_account.Credentials.from_service_account_file(SIGNER_KEY_PATH)
        _signer_creds = signer_credentials
    except Exception as e:
        print(f"⚠️ Failed to load signer credentials from {SIGNER_KEY_PATH}: {e}")
        _signer_creds = None
else:
    _signer_creds = None

def generate_signed_url(blob_name, expires_seconds=3600, disposition_filename: str | None = None):
    """
    Return a V4 signed URL for GET (downloads).
    disposition_filename optionally sets attachment filename.
    """
    if not bucket:
        raise RuntimeError("BUCKET_NAME not set or bucket not initialized")
    blob = bucket.blob(blob_name)

    response_disposition = None
    if disposition_filename:
        response_disposition = f'attachment; filename="{os.path.basename(disposition_filename)}"'

    return blob.generate_signed_url(
        version="v4",
        expiration=timedelta(seconds=expires_seconds),
        method="GET",
        response_disposition=response_disposition,
        credentials=_signer_creds,   # 👈 use the signer credentials loaded earlier
    )

# GCS sync disabled for local mode
# sync_user_tokens_from_gcs()

GEOTARGETS_LOCAL_PATH = os.environ.get("GEOTARGETS_LOCAL_PATH", "geotargets.csv")
GEOTARGETS_GCS_PATH = os.environ.get("GEOTARGETS_GCS_PATH", "geotargets.csv")
REPORT_IMAGES_GCS_PREFIX = os.environ.get("REPORT_IMAGES_GCS_PREFIX", "default_heatmaps/")

# ensure_geotargets_csv()
# ensure_default_heatmaps()

# ----- Base google-ads config -------------------------------------------
base_config = {
    "developer_token": os.environ.get("GOOGLE_ADS_DEVELOPER_TOKEN"),
    "client_id": os.environ.get("GOOGLE_ADS_CLIENT_ID"),
    "client_secret": os.environ.get("GOOGLE_ADS_CLIENT_SECRET"),
    "use_proto_plus": True,
}
if not all(base_config.values()):
    base_yaml_path = resource_path("base_google-ads.yaml")
    if os.path.exists(base_yaml_path):
        with open(base_yaml_path, "r", encoding="utf-8") as f:
            base_config.update(yaml.safe_load(f) or {})

# ----- Fixed scopes -----------------------------------------------------
SCOPES = [
    "https://www.googleapis.com/auth/adwords",
    "openid",
    "https://www.googleapis.com/auth/userinfo.email",
]
REDIRECT_URI = os.environ.get("REDIRECT_URI", None)  # MUST be set in Cloud Run env to match OAuth client

# ----- MCC Options (optional mapping for known emails) -------------------
MCC_OPTIONS = {
    "campaigns@digital.intelegencia.com": "9323527146",
    "campaigns@unyscape.com": "2626812680",
}

# ----- User token utils -------------------------------------------------
def sanitize_email_for_filename(email: str) -> str:
    return email.lower()

def user_yaml_path_for_email(email: str) -> str:
    filename = f"{sanitize_email_for_filename(email)}.yaml"
    return os.path.join(USER_TOKENS_DIR, filename)

def load_persisted_users():
    users = {}
    if not os.path.exists(USER_TOKENS_DIR):
        return users
    for fname in os.listdir(USER_TOKENS_DIR):
        if fname.endswith((".yaml", ".yml")):
            users[fname.split(".")[0]] = os.path.join(USER_TOKENS_DIR, fname)
    return users

PERSISTED_USERS = load_persisted_users()

def read_yaml(path):
    try:
        if not path or not os.path.exists(path):
            return {}
        with open(path, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception as e:
        print(f"[WARN] read_yaml failed for {path}: {e}")
        return {}

def write_yaml(path, data):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        yaml.safe_dump(data, f, sort_keys=False)
    try:
        os.chmod(path, 0o600)
    except Exception:
        pass
    # GCS persistence disabled
    pass


def normalize_customer_id(cid: str) -> str:
    return re.sub(r"[^0-9]", "", cid or "")

# ----- MCC detection & dynamic client loader -----------------------------
def is_mcc_account(client, customer_id: str) -> bool:
    """
    Return True if the given customer_id corresponds to a manager (MCC) account.
    """
    try:
        ga_service = client.get_service("GoogleAdsService")
        query = "SELECT customer.manager FROM customer LIMIT 1"
        response = ga_service.search(customer_id=customer_id, query=query)
        for row in response:
            return bool(row.customer.manager)
    except Exception as e:
        print(f"[WARN] Could not check MCC status for {customer_id}: {e}")
    return False


def find_accessible_manager(client, root_manager_id: str, target_cid: str) -> str | None:
    """
    Recursively search from root_manager_id down to find the manager that directly manages target_cid.
    Returns the accessible manager_id if found, else None.
    """
    ga_service = client.get_service("GoogleAdsService")
    query = """
        SELECT
          customer_client.id,
          customer_client.descriptive_name,
          customer_client.manager
        FROM customer_client
        WHERE customer_client.status = 'ENABLED'
    """
    try:
        response = ga_service.search(customer_id=root_manager_id, query=query)
        for row in response:
            cid = str(row.customer_client.id)
            if cid == target_cid:
                return root_manager_id  # found manager that can access target
            if row.customer_client.manager:
                # recurse down into child MCC
                mgr_id = find_accessible_manager(client, cid, target_cid)
                if mgr_id:
                    return mgr_id
    except Exception as e:
        print(f"[WARN] Error traversing from {root_manager_id} for {target_cid}: {e}")
    return None


def load_client_with_optional_login(auth_file: str, customer_id: str | None = None, manager_override: str | None = None):
    """
    Load GoogleAdsClient with correct login_customer_id handling.
    Supports:
      - Root MCC accounts
      - Nested MCCs (recursively finds parent MCC)
      - Regular client accounts
      - Optional manager overrides
    """
    cfg = read_yaml(auth_file) or {}
    cfg.setdefault("developer_token", base_config.get("developer_token"))
    cfg.setdefault("client_id", base_config.get("client_id"))
    cfg.setdefault("client_secret", base_config.get("client_secret"))
    cfg.setdefault("use_proto_plus", True)
    if "refresh_token" in cfg:
        cfg.setdefault("refresh_token", cfg.get("refresh_token"))

    from google.ads.googleads.client import GoogleAdsClient
    cid = normalize_customer_id(customer_id) if customer_id else None
    print(f"[DEBUG] load_client_with_optional_login: auth_file={auth_file}, cfg keys={list(cfg.keys())}")
    base_client = GoogleAdsClient.load_from_dict(cfg)

    if not cid:
        return base_client

    # If MCC, just return
    try:
        if is_mcc_account(base_client, cid):
            cfg.pop("login_customer_id", None)
            print(f"[INFO] {cid} is MCC -> using it directly.")
            return GoogleAdsClient.load_from_dict(cfg)
    except Exception as e:
        print(f"[WARN] MCC check failed for {cid}: {e}")

    # If client account -> determine correct MCC
    if manager_override:
        cfg["login_customer_id"] = normalize_customer_id(manager_override)
        print(f"[INFO] {cid} client -> using manager override {cfg['login_customer_id']}.")
    else:
        active_user = session.get("active_user")
        mapped_root = None
        if active_user and active_user in MCC_OPTIONS:
            mapped_root = normalize_customer_id(MCC_OPTIONS[active_user])
        if mapped_root:
            mgr_id = find_accessible_manager(base_client, mapped_root, cid)
            if mgr_id:
                cfg["login_customer_id"] = mgr_id
                print(f"[INFO] {cid} client -> resolved manager {mgr_id}.")
            else:
                print(f"[WARN] Could not resolve manager for {cid} under root {mapped_root}.")
        else:
            print(f"[WARN] No MCC mapping found for {cid}. Trying without login_customer_id.")

    return GoogleAdsClient.load_from_dict(cfg)

PERSISTED_USERS = load_persisted_users()

# ----- client secrets helper --------------------------------------------
CLIENT_SECRETS_PATH = os.environ.get("CLIENT_SECRETS_PATH", "/tmp/client-secrets-web.json")

def ensure_client_secrets_file():
    """
    Ensure we have a client secrets JSON file available.
    Priority:
     - If CLIENT_SECRETS_PATH exists, return it.
     - Else if CLIENT_SECRETS_JSON env var present, write it to CLIENT_SECRETS_PATH and return that.
     - Else check for a file named client-secrets-web.json next to this file.
    """
    if os.path.exists(CLIENT_SECRETS_PATH):
        return CLIENT_SECRETS_PATH
    cs_json = os.environ.get("CLIENT_SECRETS_JSON")
    if cs_json:
        with open(CLIENT_SECRETS_PATH, "w", encoding="utf-8") as f:
            f.write(cs_json)
        try:
            os.chmod(CLIENT_SECRETS_PATH, 0o600)
        except Exception:
            pass
        return CLIENT_SECRETS_PATH
    candidate = resource_path("client-secrets-web.json")
    if os.path.exists(candidate):
        return candidate
    raise FileNotFoundError("client-secrets-web.json not found. Set CLIENT_SECRETS_JSON or mount the file.")

# ----- Small file helpers -----------------------------------------------
def safe_remove(path):
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except Exception:
        pass


def path_in_uploads(basename: str) -> str:
    return os.path.join(app.config["UPLOAD_FOLDER"], secure_filename(os.path.basename(basename)))


# ----- Fixed routes -----------------------------------------------------
@app.route("/", methods=["GET", "POST"])
def index():
    # fallback to PERSISTED_USERS if session empty
    authenticated_users = session.get("authenticated_users", {}).copy()
    if not authenticated_users:
        authenticated_users = PERSISTED_USERS.copy()

    auth_success = session.pop("auth_success", False)
    active_user = session.get("active_user")

    if request.method == "POST":
        if not active_user or active_user not in authenticated_users:
            return redirect(url_for("auth"))

        customer_id = normalize_customer_id(request.form.get("customer_id", ""))
        # optional manager override field in the form (left blank normally)
        manager_id_override = normalize_customer_id(request.form.get("manager_id", "")) or None

        auth_file = authenticated_users[active_user]
        client = load_client_with_optional_login(auth_file, customer_id, manager_id_override)
        if client is None:
            return redirect(url_for("auth"))  # force re-auth

        # generate report - the generator should return a path to a local .docx
        filepath = generate_google_ads_report(
            customer_id=customer_id,
            google_ads_client=client,
            date_range="LAST_30_DAYS",
            start_date=None,
            end_date=None,
        )

        # normalize: ensure we have an absolute path
        if not os.path.isabs(filepath):
            filepath = os.path.abspath(filepath)

        # ensure file exists
        if not os.path.exists(filepath):
            return f"Report generator did not produce a file at expected path: {filepath}", 500

        filename = os.path.basename(filepath)
        blob_name = f"generated_reports/{filename}"

        # Upload the main report to GCS if bucket configured
        images_prefix = None
        if bucket:
            try:
                # upload_file_to_gcs is expected to upload the local filepath into the bucket and return object name/uri
                upload_file_to_gcs(filepath, blob_name)
                session["latest_report_blob"] = blob_name
            except Exception as e:
                print(f"[WARN] Failed to upload report to GCS: {e}")
        # Try to upload images (defensive: function may expect either a folder or docx path)
        try:
            # If your helper expects a directory, try candidate dirs first
            if os.path.isdir(filepath):
                images_prefix = upload_report_images_to_gcs(filepath)
            else:
                candidate_dir = os.path.splitext(filepath)[0] + "_images"
                if os.path.isdir(candidate_dir):
                    images_prefix = upload_report_images_to_gcs(candidate_dir)
                else:
                    # fallback to REPORT_IMAGES_DIR if that contains images for the report
                    if os.path.isdir(REPORT_IMAGES_DIR):
                        images_prefix = upload_report_images_to_gcs(REPORT_IMAGES_DIR)
            if images_prefix:
                session["latest_report_images_prefix"] = images_prefix
        except Exception as e:
            print(f"[WARN] upload_report_images_to_gcs failed: {e}")

        # Keep a local pointer as well so local flows work (useful for dev)
        # Copy or move file into UPLOAD_FOLDER for consistent local serving if desired
        try:
            dest = path_in_uploads(filename)
            if os.path.abspath(filepath) != os.path.abspath(dest):
                os.replace(filepath, dest)

            # ✅ Always store absolute path for parsing
            session["latest_report"] = dest
            session["latest_report_filename"] = filename
        except Exception:
            # fallback if move failed
            session["latest_report"] = filepath
            session["latest_report_filename"] = os.path.basename(filepath)

        return redirect(url_for("report"))

    return render_template(
        "index.html",
        authenticated=bool(active_user),
        auth_success=auth_success,
        authenticated_users=authenticated_users,
        active_user=active_user,
    )

@app.route("/auth")
def auth():
    session.pop("oauth_state", None)
    client_secrets = ensure_client_secrets_file()
    flow = Flow.from_client_secrets_file(client_secrets, scopes=SCOPES, redirect_uri=REDIRECT_URI)
    auth_url, state = flow.authorization_url(access_type="offline", include_granted_scopes=False, prompt="consent")
    session["oauth_state"] = state
    # PKCE: Store code_verifier to restore it in /callback
    if hasattr(flow, "code_verifier"):
        session["code_verifier"] = flow.code_verifier
    return redirect(auth_url)

@app.route("/callback")
def callback():
    stored_state = session.get("oauth_state")
    returned_state = request.args.get("state")
    if stored_state != returned_state:
        return "[WARN] OAuth state mismatch. Try <a href='/auth'>again</a>."

    client_secrets = ensure_client_secrets_file()
    flow = Flow.from_client_secrets_file(client_secrets, scopes=SCOPES, state=stored_state, redirect_uri=REDIRECT_URI)
    # Restore code_verifier for PKCE
    if "code_verifier" in session:
        flow.code_verifier = session.get("code_verifier")
    flow.fetch_token(authorization_response=request.url)
    credentials = flow.credentials

    # Verify that the adwords scope was actually granted
    if "https://www.googleapis.com/auth/adwords" not in (credentials.scopes or []):
        return "<h3>[ERROR] Permission Denied</h3><p>You must check the box to 'Manage your Google Ads campaigns' during login. <a href='/auth'>Try again</a>.</p>"

    try:
        id_info = id_token.verify_oauth2_token(credentials.id_token, requests.Request())
        email = id_info.get("email")
    except Exception as e:
        return f"Failed to obtain user email from id_token: {e}"

    refresh_token = credentials.refresh_token
    if not refresh_token:
        return "<h3>[ERROR] No refresh token returned</h3><p>Please revoke this app's access and try again.</p>"

    user_yaml = user_yaml_path_for_email(email)
    existing = read_yaml(user_yaml) if os.path.exists(user_yaml) else {}

    user_config = deepcopy(base_config)
    # 🚫 Do not persist login_customer_id blindly. Only store refresh_token.
    user_config["refresh_token"] = refresh_token

    write_yaml(user_yaml, user_config)
    # GCS Sync disabled for local mode
    # sync_user_tokens_from_gcs()

    session.setdefault("authenticated_users", {})
    session["authenticated_users"][email] = user_yaml
    PERSISTED_USERS[email] = user_yaml
    session["active_user"] = email
    session["auth_success"] = True
    session.pop("oauth_state", None)
    return redirect(url_for("index"))


@app.route("/switch_user/<identifier>")
def switch_user(identifier):
    authenticated_users = session.get("authenticated_users", {}).copy()
    if not authenticated_users:
        authenticated_users = PERSISTED_USERS.copy()
    if identifier in authenticated_users:
        session["active_user"] = identifier
        return redirect(url_for("index"))
    return "User not found.", 404


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


@app.route("/report")
def report():
    # Prefer GCS-backed report if present
    blob_name = session.get("latest_report_blob")
    if blob_name and bucket:
        tmp_path = f"/tmp/{os.path.basename(blob_name)}"
        try:
            download_gcs_to_local(blob_name, tmp_path)
            structured_report = parse_docx_to_structured(tmp_path)
        except Exception as e:
            safe_remove(tmp_path)
            return f"Failed to download or parse report from GCS: {e}", 500

        # signed URL for download + signed URLs for heatmap images
        download_link = generate_signed_url(blob_name, expires_seconds=3600, disposition_filename=os.path.basename(blob_name))
        images_prefix = session.get("latest_report_images_prefix", blob_name.rsplit(".", 1)[0] + "_images/")
        heatmap_images = []
        if bucket and images_prefix:
            for b in bucket.list_blobs(prefix=images_prefix):
                if b.name.endswith(".png"):
                    heatmap_images.append({"filename": os.path.basename(b.name), "url": generate_signed_url(b.name, expires_seconds=3600)})
        safe_remove(tmp_path)
        return render_template("report.html", structured_report=structured_report, download_link=download_link, heatmap_images=heatmap_images)

    # Fallback: local file saved in UPLOAD_FOLDER or absolute path stored in session
    filepath = session.get("latest_report")
    if filepath:
        # absolute path case
        if os.path.isabs(filepath) and os.path.exists(filepath):
            try:
                structured_report = parse_docx_to_structured(filepath)
            except Exception as e:
                return f"Failed to parse report file: {e}", 500
            download_link = url_for("download_file", filename=os.path.basename(filepath))
            return render_template("report.html", structured_report=structured_report, download_link=download_link)

        # basename in UPLOAD_FOLDER
        local_full = path_in_uploads(filepath)
        if os.path.exists(local_full):
            try:
                structured_report = parse_docx_to_structured(local_full)
            except Exception as e:
                return f"Failed to parse report file: {e}", 500
            download_link = url_for("download_file", filename=os.path.basename(local_full))
            return render_template("report.html", structured_report=structured_report, download_link=download_link)

    return "No report found. Please generate again."


@app.route("/download/<path:filename>")
def download_file(filename):
    """
    Support both:
      - GCS-backed reports (redirect to signed URL), and
      - local files in app.config['UPLOAD_FOLDER'] (send_from_directory).
    """
    # If session has a blob for the current report, and it matches this filename, redirect to signed URL
    blob_name = session.get("latest_report_blob")
    if blob_name and os.path.basename(blob_name) == os.path.basename(filename) and bucket:
        signed = generate_signed_url(blob_name, expires_seconds=3600, disposition_filename=os.path.basename(blob_name))
        return redirect(signed)

    # Fallback: serve local file from UPLOAD_FOLDER
    safe_name = secure_filename(os.path.basename(filename))
    local_path = os.path.join(app.config["UPLOAD_FOLDER"], safe_name)
    if not os.path.exists(local_path):
        return f"File not found on server: {local_path}", 404
    return send_from_directory(app.config["UPLOAD_FOLDER"], safe_name, as_attachment=True)

@app.route("/download/")
def download_latest():
    """
    Fallback route so /download/ (no filename) redirects to the latest generated report.
    Uses session['latest_report_filename'] if present, else falls back to the basename of session['latest_report'].
    """
    filename = session.get("latest_report_filename")
    if not filename:
        latest = session.get("latest_report")
        if latest:
            filename = os.path.basename(latest)
    if not filename:
        return "No report available for download.", 404
    return redirect(url_for("download_file", filename=filename))

@app.route("/section/<section_id>")
def section_detail(section_id):
    filepath = session.get("latest_report")
    if not filepath or not os.path.exists(path_in_uploads(filepath)):
        # allow GCS-backed heatmaps through session images prefix when appropriate
        pass

    if section_id == "heatmaps":
        charts = []
        images_prefix = session.get("latest_report_images_prefix")
        if images_prefix and bucket:
            for blob in bucket.list_blobs(prefix=images_prefix):
                if blob.name.endswith(".png"):
                    parts = os.path.splitext(os.path.basename(blob.name))[0].split("_")
                    charts.append(
                        {
                            "campaign_type": parts[0] if len(parts) >= 2 else "General",
                            "metric": parts[1] if len(parts) >= 2 else parts[0],
                            "filename": os.path.basename(blob.name),
                            "url": generate_signed_url(blob.name),
                        }
                    )
        else:
            for fname in os.listdir(REPORT_IMAGES_DIR):
                if fname.endswith(".png"):
                    parts = os.path.splitext(fname)[0].split("_")
                    charts.append(
                        {
                            "campaign_type": parts[0] if len(parts) >= 2 else "General",
                            "metric": parts[1] if len(parts) >= 2 else parts[0],
                            "filename": fname,
                            "url": url_for("report_images", filename=fname),
                        }
                    )
        return render_template("section.html", section={"title": "📊 Heatmaps", "content": [{"type": "images", "charts": charts}]})

    try:
        section_index = int(section_id)
        # if blob-backed, we already handle heatmaps; for section text we need a local docx
        local_path = None
        blob_name = session.get("latest_report_blob")
        if blob_name and bucket:
            tmp_path = f"/tmp/{os.path.basename(blob_name)}"
            try:
                download_gcs_to_local(blob_name, tmp_path)
                structured_report = parse_docx_to_structured(tmp_path)
            finally:
                safe_remove(tmp_path)
        else:
            filepath = session.get("latest_report")
            local_path = path_in_uploads(filepath) if filepath else None
            if not local_path or not os.path.exists(local_path):
                return "No report found. Please generate again."
            structured_report = parse_docx_to_structured(local_path)

        if section_index < 0 or section_index >= len(structured_report):
            return "Invalid section ID."
        return render_template("section.html", section=structured_report[section_index])
    except ValueError:
        return "Invalid section ID (not a number)."


@app.route("/report_images/<filename>")
def report_images(filename):
    return send_from_directory(REPORT_IMAGES_DIR, filename)


# ----- Table/docx parsing utils -------------------------------------------
def try_parse_to_table(text):
    if not text or not isinstance(text, str):
        return None
    cleaned = text.strip().replace("“", '"').replace("”", '"')
    try:
        data = json.loads(cleaned)
        if isinstance(data, dict):
            data = [data]
        if isinstance(data, list) and all(isinstance(d, dict) for d in data):
            headers = list(data[0].keys())
            rows = [[d.get(h, "") for h in headers] for d in data]
            return {"headers": headers, "rows": rows}
    except Exception:
        try:
            data = ast.literal_eval(cleaned)
            if isinstance(data, list) and all(isinstance(d, dict) for d in data):
                headers = list(data[0].keys())
                rows = [[d.get(h, "") for h in headers] for d in data]
                return {"headers": headers, "rows": rows}
        except Exception:
            pass

    lines = [line.strip("•*- ") for line in cleaned.splitlines() if "|" in line]
    rows = [line.split("|")[:3] for line in lines if len(line.split("|")) >= 3]
    if rows:
        return {"headers": ["Characteristic", "Insight", "Recommendation"], "rows": rows}

    for sep in [",", "\t"]:
        lines = [line for line in cleaned.splitlines() if sep in line]
        rows = [line.split(sep)[:3] for line in lines if len(line.split(sep)) >= 3]
        if rows:
            return {"headers": ["Characteristic", "Insight", "Recommendation"], "rows": rows}
    return None


def parse_docx_to_structured(path):
    from docx import Document

    doc = Document(path)
    structured = []
    current_section = {
        "title": "⭐ Introduction",
        "content": [
            {
                "type": "paragraph",
                "content": (
                    "Hello, this is a structured report generated from a Google Ads audit document. "
                    "It contains insights, visualizations, and optimization suggestions."
                ),
            }
        ],
    }
    para_index = 0
    table_index = 0

    for element in doc.element.body:
        if element.tag.endswith("}p"):
            if para_index >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[para_index]
            para_index += 1
            text = para.text.strip()
            if not text:
                continue
            if getattr(para.style, "name", "").startswith("Heading") or text.startswith("⭐"):
                if current_section["content"]:
                    structured.append(current_section)
                current_section = {"title": text, "content": []}
            else:
                table_result = try_parse_to_table(text)
                if table_result:
                    current_section["content"].append(
                        {"type": "table", "headers": table_result["headers"], "rows": table_result["rows"]}
                    )
                else:
                    current_section["content"].append({"type": "paragraph", "content": text})
        elif element.tag.endswith("}tbl"):
            if table_index >= len(doc.tables):
                continue
            table = doc.tables[table_index]
            table_index += 1
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                current_section["content"].append({"type": "table", "headers": rows[0], "rows": rows[1:]})

    if current_section["content"]:
        structured.append(current_section)
    return structured

if __name__ == "__main__":
    PORT = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=PORT, debug=True, use_reloader=False)